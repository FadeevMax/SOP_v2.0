import streamlit as st
from openai import OpenAI
import time
import os
import uuid
from datetime import datetime, timedelta
import json
from streamlit_local_storage import LocalStorage
import difflib
# Imports for Google Docs API
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
import io # Needed for handling the in-memory file download
import requests
import base64
import unicodedata
from utils.config import GDOC_STATE_PATH, GOOGLE_DOC_NAME, CACHE_DIR, PDF_CACHE_PATH, DOCX_LOCAL_PATH, IMAGE_DIR, IMAGE_MAP_PATH, ENRICHED_CHUNKS_PATH, GITHUB_REPO, GITHUB_TOKEN
import re
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from utils.github import update_pdf_on_github, update_docx_on_github, update_json_on_github, upload_file_to_github
import numpy as np
from sentence_transformers import SentenceTransformer

caption_pattern = re.compile(r"^Image\s+(\d+):?\s*(.*)", re.IGNORECASE)

def clean_caption(text):
    cleaned = unicodedata.normalize('NFKC', text)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    cleaned = cleaned.replace("–", "-").replace("—", "-").replace(""", '"').replace(""", '"')
    cleaned = cleaned.replace("'", "'").replace("'", "'")
    return cleaned

def extract_label(text):
    text = clean_caption(text)
    m = caption_pattern.match(text)
    if m:
        idx = int(m.group(1))
        desc = m.group(2).strip().rstrip(".")
        return f"Image {idx}: {desc}" if desc else f"Image {idx}"
    return None

def extract_images_and_labels_from_docx(docx_path, image_output_dir, mapping_output_path, debug=False):
    """Extract images and their labels from a DOCX file"""
    os.makedirs(image_output_dir, exist_ok=True)
    doc = Document(docx_path)
    image_map = {}
    items = []
    
    # Collect all blocks in order
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            # Check for images in paragraph
            has_image = False
            for run in para.runs:
                if 'graphic' in run._element.xml:
                    for drawing in run._element.findall(".//w:drawing", namespaces=run._element.nsmap):
                        for blip in drawing.findall(".//a:blip", namespaces=run._element.nsmap):
                            rel_id = blip.get(qn('r:embed'))
                            if rel_id and rel_id in doc.part.related_parts:
                                image_part = doc.part.related_parts[rel_id]
                                items.append(('image', image_part))
                                has_image = True
            
            # Add text if it exists
            if para.text.strip():
                items.append(('text', para.text.strip()))
                
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Check for images in table cells
                        for run in para.runs:
                            if 'graphic' in run._element.xml:
                                for drawing in run._element.findall(".//w:drawing", namespaces=run._element.nsmap):
                                    for blip in drawing.findall(".//a:blip", namespaces=run._element.nsmap):
                                        rel_id = blip.get(qn('r:embed'))
                                        if rel_id and rel_id in doc.part.related_parts:
                                            image_part = doc.part.related_parts[rel_id]
                                            items.append(('image', image_part))
                        
                        if para.text.strip():
                            items.append(('text', para.text.strip()))

    # Associate images with their following captions
    image_counter = 1
    i = 0
    while i < len(items):
        if items[i][0] == 'image':
            image_part = items[i][1]
            
            # Look for the next text that might be a caption
            label = None
            for j in range(i + 1, min(i + 3, len(items))):  # Look ahead up to 2 items
                if items[j][0] == 'text':
                    potential_label = extract_label(items[j][1])
                    if potential_label:
                        label = potential_label
                        break
            
            if not label:
                label = f"Image {image_counter}"
            
            # Save image file
            image_extension = image_part.content_type.split('/')[-1]
            if image_extension == 'jpeg':
                image_extension = 'jpg'
            image_name = f"image_{image_counter}.{image_extension}"
            image_path = os.path.join(image_output_dir, image_name)
            
            with open(image_path, "wb") as f:
                f.write(image_part.blob)
            
            image_map[label] = image_name
            image_counter += 1
            
        i += 1

    # Save mapping
    with open(mapping_output_path, "w") as f:
        json.dump(image_map, f, indent=2)
    
    if debug:
        print("Final image_map:")
        for caption, img in image_map.items():
            print(f"{caption} => {img}")
    
    return image_map

def force_resync_to_github():
    """
    Forces the re-processing of the local DOCX file and syncs all assets to GitHub.
    This skips the Google Doc check and works with the current local DOCX.
    """
    if not os.path.exists(DOCX_LOCAL_PATH):
        st.error("Cannot re-sync: The local DOCX file does not exist.")
        return False

    try:
        # Step 1: Re-extract images and create map.json from the local DOCX
        st.write("Extracting images and labels from local DOCX...")
        extract_images_and_labels_from_docx(DOCX_LOCAL_PATH, IMAGE_DIR, IMAGE_MAP_PATH, debug=True)
        st.write("✅ Image extraction complete.")

        # Step 2: Upload map.json to GitHub
        st.write("Uploading map.json to GitHub...")
        update_json_on_github(
            local_json_path=IMAGE_MAP_PATH,
            repo_json_path="map.json",
            commit_message="Manual Re-sync: Update map.json",
            github_repo=GITHUB_REPO,
            github_token=GITHUB_TOKEN
        )
        st.write("✅ map.json uploaded.")

        # Step 3: Upload all images from the local image directory
        st.write("Uploading images to GitHub...")
        for file in os.listdir(IMAGE_DIR):
            local_path = os.path.join(IMAGE_DIR, file)
            github_path = f"images/{file}"
            upload_file_to_github(
                local_path=local_path,
                github_path=github_path,
                commit_message=f"Manual Re-sync: Update {file}"
            )
        st.write("✅ Images uploaded.")

        # Step 3.5: Generate enriched content chunks and upload
        st.write("Generating semantic content chunks...")
        chunks = semantic_chunking_docx(DOCX_LOCAL_PATH)
        enriched_chunks = enrich_chunks_with_images_semantic(chunks, IMAGE_MAP_PATH)
        with open(ENRICHED_CHUNKS_PATH, "w") as f:
            json.dump(enriched_chunks, f, indent=2)
        # Also write the combined chunk text file
        chunk_text_path = os.path.join(CACHE_DIR, "sop_chunks.txt")
        with open(chunk_text_path, "w") as f:
            for chunk in enriched_chunks:
                f.write(chunk["chunk_text"] + "\n\n")
        st.write("✅ Enriched chunks file created.")
        # Upload enriched_chunks.json to GitHub
        st.write("Uploading enriched_chunks.json to GitHub...")
        upload_file_to_github(
            local_path=ENRICHED_CHUNKS_PATH,
            github_path="enriched_chunks.json",
            commit_message="Manual Re-sync: Update enriched_chunks.json"
        )
        st.write("✅ enriched_chunks.json uploaded.")

        # Step 4: Upload the DOCX and PDF files
        st.write("Uploading DOCX and PDF to GitHub...")
        update_docx_on_github(DOCX_LOCAL_PATH)
        if os.path.exists(PDF_CACHE_PATH):
             update_pdf_on_github(PDF_CACHE_PATH)
        st.write("✅ Document files uploaded.")

        return True

    except Exception as e:
        st.error(f"An error occurred during the re-sync process: {e}")
        return False
        
def get_creds():
    """Get credentials from Streamlit secrets or local JSON file."""
    try:
        # Try to get from Streamlit secrets first
        creds_dict = st.secrets["gcp_service_account"]
        return Credentials.from_service_account_info(creds_dict)
    except (KeyError, FileNotFoundError):
        # If not found, try to use a local file
        local_path = "gcp_service_account.json"
        if os.path.exists(local_path):
            return Credentials.from_service_account_file(local_path)
        else:
            st.error("GCP service account credentials not found.")
            st.stop()

def download_gdoc_as_docx(doc_id, creds, out_path):
   drive_service = build('drive', 'v3', credentials=creds)
   request = drive_service.files().export_media(fileId=doc_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
   os.makedirs(os.path.dirname(out_path), exist_ok=True)
   with open(out_path, "wb") as f:
     f.write(request.execute())
   return True

def download_gdoc_as_pdf(doc_id, creds, out_path):
    drive_service = build('drive', 'v3', credentials=creds)
    request = drive_service.files().export_media(fileId=doc_id, mimeType='application/pdf')
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "wb") as f:
        f.write(request.execute())
    return True

def get_gdoc_last_modified(creds, doc_name):
    drive_service = build('drive', 'v3', credentials=creds)
    query = f"name='{doc_name}' and mimeType='application/vnd.google-apps.document'"
    results = drive_service.files().list(q=query, fields="files(id, modifiedTime)").execute()
    files = results.get('files', [])
    if not files:
        return None, None
    doc_id = files[0]['id']
    modified_time = files[0]['modifiedTime']
    return doc_id, modified_time

def get_live_sop_pdf_path(doc_name: str) -> str:
    """
    Checks for a fresh cached PDF. If it's stale or missing, it downloads
    the live Google Doc as a PDF and saves it to the cache.
    Returns the file path to the fresh PDF.
    """
    try:
        st.info("Checking for SOP updates from Google Docs...")

        # Create cache directory if it doesn't exist
        if not os.path.exists(CACHE_DIR):
            os.makedirs(CACHE_DIR)

        creds = get_creds()
        drive_service = build('drive', 'v3', credentials=creds)

        query = f"name='{doc_name}' and mimeType='application/vnd.google-apps.document'"
        response = drive_service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
        files = response.get('files', [])

        if not files:
            st.error(f"No Google Doc found with the name: '{doc_name}'.")
            return None

        doc_id = files[0]['id']

        # Export the document as PDF
        request = drive_service.files().export_media(fileId=doc_id, mimeType='application/pdf')

        # Download the file content into an in-memory buffer
        fh = io.BytesIO()
        downloader = io.BytesIO(request.execute())

        # Define the path for the cached file
        cached_file_path = os.path.join(CACHE_DIR, "cached_sop.pdf")

        # Write the downloaded content to the cached file
        with open(cached_file_path, "wb") as f:
            f.write(downloader.getbuffer())

        st.success(f"✅ SOP updated successfully from Google Docs!")
        return cached_file_path

    except Exception as e:
        st.error(f"❌ Failed to fetch and cache Google Doc: {e}")
        return None


def get_last_gdoc_synced_time():
    if os.path.exists(GDOC_STATE_PATH):
        with open(GDOC_STATE_PATH, "r") as f:
            state = json.load(f)
            return state.get("last_synced_modified_time")
    return None

def set_last_gdoc_synced_time(modified_time):
    with open(GDOC_STATE_PATH, "w") as f:
        json.dump({"last_synced_modified_time": modified_time}, f)

def sync_gdoc_to_github(force=False):
    # Only check if a day has passed or force=True
    last_synced = get_last_gdoc_synced_time()
    now = datetime.utcnow()
    last_checked_dt = datetime.fromisoformat(last_synced) if last_synced else None

    # Google API Auth
    creds = get_creds()
    doc_id, modified_time = get_gdoc_last_modified(creds, GOOGLE_DOC_NAME)
    if not doc_id or not modified_time:
        st.warning("Google Doc not found or can't fetch modified time.")
        return False

    # Only update if new or forced or more than 1 day has passed
    need_update = (
        force or 
        not last_synced or 
        (now - last_checked_dt > timedelta(days=1)) or
        (modified_time != last_synced)
    )
    if not need_update:
        st.info("No update needed. Using existing GitHub PDF.")
        return True

    # Download latest Google Doc as PDF and DOCX
    pdf_success = download_gdoc_as_pdf(doc_id, creds, PDF_CACHE_PATH)
    docx_success = download_gdoc_as_docx(doc_id, creds, DOCX_LOCAL_PATH)

    if not (pdf_success and docx_success):
       st.error("Failed to download Google Doc as PDF or DOCX.")
       return False

    # Extract labeled images from DOCX
    extract_images_and_labels_from_docx(DOCX_LOCAL_PATH, IMAGE_DIR, IMAGE_MAP_PATH, debug=True)

    # Update map.json on GitHub
    success = update_json_on_github(
       IMAGE_MAP_PATH,
       "map.json",
       "Update map.json from SOP DOCX",
       GITHUB_REPO,
       GITHUB_TOKEN
   )
    if not success:
       st.error("❌ Failed to update map.json on GitHub!")

    # Upload images to GitHub
    for file in os.listdir(IMAGE_DIR):
       local_path = os.path.join(IMAGE_DIR, file)
       github_path = f"images/{file}"
       upload_file_to_github(
          local_path,
          github_path,
          f"Update {file} from SOP DOCX"
          )

    # Generate semantic chunks from the DOCX and save to cache
    st.info("Generating semantic chunks of the SOP content...")
    chunks = semantic_chunking_docx(DOCX_LOCAL_PATH)
    enriched_chunks = enrich_chunks_with_images_semantic(chunks, IMAGE_MAP_PATH)
    # Save enriched chunks JSON
    with open(ENRICHED_CHUNKS_PATH, "w") as f:
        json.dump(enriched_chunks, f, indent=2)
    # Save combined chunk text to a .txt file for vector store use
    chunk_text_path = os.path.join(CACHE_DIR, "sop_chunks.txt")
    with open(chunk_text_path, "w") as f:
        for chunk in enriched_chunks:
            f.write(chunk["chunk_text"] + "\n\n")
    st.success(f"✅ Split SOP into {len(enriched_chunks)} enriched chunks.")

    # Upload PDF and DOCX to GitHub
    pdf_uploaded = update_pdf_on_github(PDF_CACHE_PATH)
    docx_uploaded = update_docx_on_github(DOCX_LOCAL_PATH)
    # Upload enriched_chunks.json if it exists
    if os.path.exists(ENRICHED_CHUNKS_PATH):
        upload_file_to_github(
            local_path=ENRICHED_CHUNKS_PATH,
            github_path="enriched_chunks.json",
            commit_message="Update enriched chunks"
        )
    else:
        st.warning(f"enriched_chunks.json not found at {ENRICHED_CHUNKS_PATH}, skipping upload.")

    if pdf_uploaded and docx_uploaded:
        st.success("PDF and DOCX updated on GitHub with the latest from Google Doc!")
        set_last_gdoc_synced_time(modified_time)
        return True
    elif pdf_uploaded:
        st.error("PDF uploaded, but failed to update DOCX on GitHub.")
        return False
    elif docx_uploaded:
        st.error("DOCX uploaded, but failed to update PDF on GitHub.")
        return False
    else:
        st.error("Failed to update both PDF and DOCX on GitHub.")
        return False

# --- Semantic Chunking Utilities ---
def semantic_chunking_docx(docx_path, model_name='all-MiniLM-L6-v2', buffer_size=1, percentile=90, min_chunk_size=50):
    """
    Semantic chunking of DOCX content with image association.
    Image captions are always grouped with the text above.
    """
    doc = Document(docx_path)
    model = SentenceTransformer(model_name)

    # Extract all content paragraphs with metadata
    content_items = []
    prev_text_idx = None
    for para in doc.paragraphs:
        text = clean_caption(para.text.strip())
        if not text:
            continue
        is_caption = bool(caption_pattern.match(text))
        if is_caption and prev_text_idx is not None:
            # If this paragraph is an image caption, append it to the previous text chunk
            content_items[prev_text_idx]['text'] += f"\n{text}"
            content_items[prev_text_idx]['is_caption'] = True
            continue
        else:
            content_items.append({
                'text': text,
                'is_caption': is_caption,
                'paragraph_index': len(content_items)
            })
            prev_text_idx = len(content_items) - 1

    # Filter out very short items (unless they contain a caption)
    sentences = []
    for i, item in enumerate(content_items):
        if len(item['text']) >= min_chunk_size or item['is_caption']:
            sentences.append({
                'sentence': item['text'],
                'is_caption': item['is_caption'],
                'original_index': i
            })
    if not sentences:
        return []

    # Create contextual buffer for each sentence (previous and next sentence context)
    for i in range(len(sentences)):
        context = []
        for j in range(i - buffer_size, i + buffer_size + 1):
            if 0 <= j < len(sentences):
                context.append(sentences[j]['sentence'])
        sentences[i]['combined'] = ' '.join(context)

    # Generate embeddings for each combined context
    embeddings = model.encode([s['combined'] for s in sentences], convert_to_numpy=True)
    for i, emb in enumerate(embeddings):
        sentences[i]['embedding'] = emb

    # Compute cosine distance between consecutive sentence embeddings
    distances = []
    for i in range(len(sentences) - 1):
        a, b = sentences[i]['embedding'], sentences[i+1]['embedding']
        dist = 1 - (np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b)))
        distances.append(dist)
        sentences[i]['distance_to_next'] = dist

    # Determine cut points where distance exceeds the percentile threshold
    threshold = np.percentile(distances, percentile) if distances else float('inf')
    breakpoints = [i for i, d in enumerate(distances) if d > threshold]

    # Assemble chunks based on breakpoints
    chunks = []
    start_idx = 0
    for bp in breakpoints:
        end_idx = bp + 1
        chunk_sentences = sentences[start_idx:end_idx]
        if chunk_sentences:
            chunks.append(chunk_sentences)
        start_idx = end_idx
    if start_idx < len(sentences):
        chunks.append(sentences[start_idx:])

    # If no semantic breaks found, fall back to fixed-size chunking (to avoid one huge chunk)
    if not chunks and sentences:
        chunk_size = max(3, len(sentences) // 5)  # aim for ~5 chunks if possible
        chunks = [sentences[i:i + chunk_size] for i in range(0, len(sentences), chunk_size)]

    return chunks

def enrich_chunks_with_images_semantic(chunks, image_map_path):
    """
    Attach image labels/files to each chunk based on references in the text.
    Only includes images that have labels present in the chunk text.
    """
    try:
        with open(image_map_path, "r") as f:
            image_map = json.load(f)
    except FileNotFoundError:
        image_map = {}

    def extract_image_numbers(text):
        patterns = [r"Image\s+(\d+)", r"Figure\s+(\d+)"]
        nums = []
        for pattern in patterns:
            nums.extend(re.findall(pattern, text, flags=re.IGNORECASE))
        return {int(n) for n in nums}

    def find_image_by_number(img_num):
        for label, filename in image_map.items():
            # If the label starts with "Image X" matching the number, return it
            if label.lower().startswith(f"image {img_num}"):
                return label, filename
        return None, None

    enriched = []
    used_labels = set()
    for idx, chunk in enumerate(chunks):
        # Combine sentences back into a chunk of text
        chunk_text = "\n".join([s['sentence'] for s in chunk])
        image_numbers = extract_image_numbers(chunk_text)
        images = []
        for num in image_numbers:
            label, fname = find_image_by_number(num)
            if label and label in image_map:
                images.append({"label": label, "file": fname})
                used_labels.add(label)
        enriched.append({
            "chunk_id": idx,
            "chunk_text": chunk_text,
            "sentence_count": len(chunk),
            "image_files": images,
            "has_images": len(images) > 0
        })
    return enriched
