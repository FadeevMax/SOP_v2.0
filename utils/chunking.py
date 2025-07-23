import streamlit as st
from openai import OpenAI
import time
import os
import uuid
from datetime import datetime
import json
from streamlit_local_storage import LocalStorage
import difflib
# Imports for Google Docs API
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
import io # Needed for handling the in-memory file download
import requests
import base64
import unicodedata
import re 
import numpy as np
from sentence_transformers import SentenceTransformer
from utils.config import (
    CACHE_DIR,
    PDF_CACHE_PATH,
    GDOC_STATE_PATH,
    GITHUB_PDF_NAME,
    GITHUB_REPO,
    GITHUB_TOKEN,
    GOOGLE_DOC_NAME,
    STATE_DIR,
    DOCX_LOCAL_PATH,
    IMAGE_DIR,
    IMAGE_MAP_PATH,
    ENRICHED_CHUNKS_PATH,
)
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from sklearn.metrics.pairwise import cosine_similarity
from scipy.signal import find_peaks

# Enhanced patterns and functions
caption_pattern = re.compile(r"^Image\s+(\d+):?\s*(.*)", re.IGNORECASE)
figure_pattern = re.compile(r"^Figure\s+(\d+):?\s*(.*)", re.IGNORECASE)
reference_pattern = re.compile(r"(?:see|refer to|shown in|as depicted in|according to)\s+(?:image|figure)\s+(\d+)", re.IGNORECASE)

def clean_caption(text):
    """Enhanced text cleaning with better normalization"""
    cleaned = unicodedata.normalize('NFKC', text)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    cleaned = cleaned.replace("–", "-").replace("—", "-").replace(""", '"').replace(""", '"')
    cleaned = cleaned.replace("'", "'").replace("'", "'")
    # Remove excessive punctuation
    cleaned = re.sub(r'[.]{2,}', '.', cleaned)
    return cleaned

def extract_label(text):
    """Enhanced label extraction supporting multiple formats"""
    text = clean_caption(text)
    
    # Try Image pattern first
    m = caption_pattern.match(text)
    if m:
        idx = int(m.group(1))
        desc = m.group(2).strip().rstrip(".")
        return f"Image {idx}: {desc}" if desc else f"Image {idx}"
    
    # Try Figure pattern
    m = figure_pattern.match(text)
    if m:
        idx = int(m.group(1))
        desc = m.group(2).strip().rstrip(".")
        return f"Figure {idx}: {desc}" if desc else f"Figure {idx}"
    
    return None

def extract_image_references(text):
    """Extract references to images/figures in text"""
    references = []
    
    # Find direct references like "see Image 3" or "shown in Figure 2"
    matches = reference_pattern.findall(text)
    references.extend([int(match) for match in matches])
    
    # Find standalone mentions like "Image 5" or "Figure 3"
    standalone_matches = re.findall(r"(?:Image|Figure)\s+(\d+)", text, re.IGNORECASE)
    references.extend([int(match) for match in standalone_matches])
    
    return list(set(references))  # Remove duplicates

def extract_images_and_labels_from_docx(docx_path, image_output_dir, mapping_output_path, debug=False):
    """Enhanced image extraction with better ordering and metadata"""
    os.makedirs(image_output_dir, exist_ok=True)
    doc = Document(docx_path)
    image_map = {}
    items = []
    
    # Collect all blocks in document order with position tracking
    body = doc.element.body
    position = 0
    
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            para = Paragraph(child, doc)
            
            # Check for images in paragraph
            for run in para.runs:
                if 'graphic' in run._element.xml:
                    for drawing in run._element.findall(".//w:drawing", namespaces=run._element.nsmap):
                        for blip in drawing.findall(".//a:blip", namespaces=run._element.nsmap):
                            rel_id = blip.get(qn('r:embed'))
                            if rel_id and rel_id in doc.part.related_parts:
                                image_part = doc.part.related_parts[rel_id]
                                items.append({
                                    'type': 'image', 
                                    'content': image_part, 
                                    'position': position,
                                    'paragraph_text': para.text.strip()
                                })
            
            # Add text if it exists
            if para.text.strip():
                items.append({
                    'type': 'text', 
                    'content': para.text.strip(), 
                    'position': position
                })
                
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Check for images in table cells
                        for run in para.runs:
                            if 'graphic' in run._element.xml:
                                for drawing in run._element.findall(".//w:drawing", namespaces=run._element.nsmap):
                                    for blip in drawing.findall(".//a:blip", namespaces=run._element.nsmap):
                                        rel_id = blip.get(qn('r:embed'))
                                        if rel_id and rel_id in doc.part.related_parts:
                                            image_part = doc.part.related_parts[rel_id]
                                            items.append({
                                                'type': 'image', 
                                                'content': image_part, 
                                                'position': position,
                                                'paragraph_text': para.text.strip()
                                            })
                        
                        if para.text.strip():
                            items.append({
                                'type': 'text', 
                                'content': para.text.strip(), 
                                'position': position
                            })
        
        position += 1

    # Enhanced image-caption association
    image_counter = 1
    i = 0
    
    while i < len(items):
        if items[i]['type'] == 'image':
            image_part = items[i]['content']
            
            # Look for caption in multiple places
            label = None
            
            # 1. Check if the image's paragraph contains caption text
            if items[i].get('paragraph_text'):
                potential_label = extract_label(items[i]['paragraph_text'])
                if potential_label:
                    label = potential_label
            
            # 2. Look ahead for following caption
            if not label:
                for j in range(i + 1, min(i + 4, len(items))):  # Look ahead up to 3 items
                    if items[j]['type'] == 'text':
                        potential_label = extract_label(items[j]['content'])
                        if potential_label:
                            label = potential_label
                            break
            
            # 3. Look behind for preceding caption
            if not label:
                for j in range(max(0, i - 3), i):  # Look behind up to 3 items
                    if items[j]['type'] == 'text':
                        potential_label = extract_label(items[j]['content'])
                        if potential_label:
                            label = potential_label
                            break
            
            # Default label if none found
            if not label:
                label = f"Image {image_counter}"
            
            # Save image file with better naming
            image_extension = image_part.content_type.split('/')[-1]
            if image_extension == 'jpeg':
                image_extension = 'jpg'
            elif image_extension not in ['jpg', 'png', 'gif', 'bmp', 'webp']:
                image_extension = 'png'  # Default fallback
                
            image_name = f"image_{image_counter}.{image_extension}"
            image_path = os.path.join(image_output_dir, image_name)
            
            with open(image_path, "wb") as f:
                f.write(image_part.blob)
            
            image_map[label] = {
                'file': image_name,
                'position': items[i]['position'],
                'counter': image_counter
            }
            image_counter += 1
            
        i += 1

    # Convert to simpler format for backwards compatibility
    simple_map = {k: v['file'] for k, v in image_map.items()}
    
    # Save both formats
    with open(mapping_output_path, "w") as f:
        json.dump(simple_map, f, indent=2)
        
    # Save detailed mapping for advanced chunking
    detailed_map_path = mapping_output_path.replace('.json', '_detailed.json')
    with open(detailed_map_path, "w") as f:
        json.dump(image_map, f, indent=2)
    
    if debug:
        print("Final image_map:")
        for caption, info in image_map.items():
            print(f"{caption} => {info['file']} (pos: {info['position']})")
    
    return image_map

# --- Advanced Semantic Chunking ---
def semantic_chunking_docx(docx_path, model_name='all-MiniLM-L6-v2', 
                                   buffer_size=2, percentile=85, min_chunk_size=100,
                                   max_chunk_size=2000, similarity_threshold=0.3):
    """
    Advanced semantic chunking with multiple strategies:
    1. Contextual buffering around sentences
    2. Multiple similarity thresholds
    3. Content-aware boundaries (respect paragraph breaks)
    4. Dynamic chunk sizing
    """
    doc = Document(docx_path)
    model = SentenceTransformer(model_name)

    # Extract content with enhanced metadata
    content_items = []
    for para_idx, para in enumerate(doc.paragraphs):
        text = clean_caption(para.text.strip())
        if not text:
            continue

        is_caption = bool(caption_pattern.match(text) or figure_pattern.match(text))
        image_refs = extract_image_references(text)
        
        content_items.append({
            'text': text,
            'is_caption': is_caption,
            'paragraph_index': para_idx,
            'char_length': len(text),
            'word_count': len(text.split()),
            'image_references': image_refs,
            'is_short': len(text) < min_chunk_size
        })

    # Filter and prepare sentences
    sentences = []
    for i, item in enumerate(content_items):
        # Keep captions and substantial text
        if item['char_length'] >= min_chunk_size or item['is_caption'] or item['image_references']:
            sentences.append({
                'sentence': item['text'],
                'is_caption': item['is_caption'],
                'original_index': i,
                'paragraph_index': item['paragraph_index'],
                'image_references': item['image_references'],
                'char_length': item['char_length']
            })

    if not sentences:
        return []

    # Enhanced contextual buffering
    for i in range(len(sentences)):
        # Create overlapping context windows
        contexts = []
        
        # Local context (immediate neighbors)
        local_combo = []
        for j in range(i - buffer_size, i + buffer_size + 1):
            if 0 <= j < len(sentences):
                local_combo.append(sentences[j]['sentence'])
        
        # Topic context (broader semantic window)
        topic_combo = []
        topic_window = min(buffer_size * 2, 5)  # Larger window for topic coherence
        for j in range(i - topic_window, i + topic_window + 1):
            if 0 <= j < len(sentences):
                topic_combo.append(sentences[j]['sentence'])
        
        sentences[i]['local_context'] = ' '.join(local_combo)
        sentences[i]['topic_context'] = ' '.join(topic_combo[:500])  # Limit length

    # Generate embeddings for multiple contexts
    if len(sentences) > 1:
        local_contexts = [s['local_context'] for s in sentences]
        topic_contexts = [s['topic_context'] for s in sentences]
        
        local_embeddings = model.encode(local_contexts, convert_to_numpy=True)
        topic_embeddings = model.encode(topic_contexts, convert_to_numpy=True)

        for i, (local_emb, topic_emb) in enumerate(zip(local_embeddings, topic_embeddings)):
            sentences[i]['local_embedding'] = local_emb
            sentences[i]['topic_embedding'] = topic_emb

        # Compute multiple distance metrics
        local_distances = []
        topic_distances = []
        combined_distances = []
        
        for i in range(len(sentences) - 1):
            # Local similarity (immediate coherence)
            local_sim = cosine_similarity([sentences[i]['local_embedding']], 
                                        [sentences[i + 1]['local_embedding']])[0][0]
            local_dist = 1 - local_sim
            
            # Topic similarity (broader coherence)
            topic_sim = cosine_similarity([sentences[i]['topic_embedding']], 
                                        [sentences[i + 1]['topic_embedding']])[0][0]
            topic_dist = 1 - topic_sim
            
            # Combined weighted distance
            combined_dist = 0.6 * local_dist + 0.4 * topic_dist
            
            local_distances.append(local_dist)
            topic_distances.append(topic_dist)
            combined_distances.append(combined_dist)
            
            sentences[i]['local_distance'] = local_dist
            sentences[i]['topic_distance'] = topic_dist
            sentences[i]['combined_distance'] = combined_dist

        # Smart breakpoint detection
        breakpoints = find_smart_breakpoints(sentences, combined_distances, percentile, 
                                           similarity_threshold, min_chunk_size, max_chunk_size)
    else:
        breakpoints = []

    # Create semantically coherent chunks
    chunks = create_smart_chunks(sentences, breakpoints, min_chunk_size, max_chunk_size)
    
    return chunks

def find_smart_breakpoints(sentences, distances, percentile, similarity_threshold, 
                          min_chunk_size, max_chunk_size):
    """Find breakpoints using multiple strategies"""
    if not distances:
        return []
    
    breakpoints = set()
    
    # Strategy 1: Statistical threshold
    threshold = np.percentile(distances, percentile)
    statistical_breaks = [i for i, d in enumerate(distances) if d > threshold]
    
    # Strategy 2: Peak detection for natural boundaries
    peaks, _ = find_peaks(distances, height=similarity_threshold, distance=2)
    
    # Strategy 3: Content-aware breaks (prefer paragraph boundaries)
    content_breaks = []
    for i, sentence in enumerate(sentences[:-1]):
        if (sentence['is_caption'] or 
            sentences[i + 1]['is_caption'] or
            abs(sentence['paragraph_index'] - sentences[i + 1]['paragraph_index']) > 1):
            content_breaks.append(i)
    
    # Combine strategies
    all_candidates = set(statistical_breaks + list(peaks) + content_breaks)
    
    # Filter based on chunk size constraints
    confirmed_breaks = []
    last_break = 0
    
    for candidate in sorted(all_candidates):
        # Check if this would create a reasonable chunk size
        chunk_chars = sum(s['char_length'] for s in sentences[last_break:candidate + 1])
        
        if chunk_chars >= min_chunk_size:
            confirmed_breaks.append(candidate)
            last_break = candidate + 1
        elif chunk_chars > max_chunk_size and candidate - last_break > 0:
            # Force break if chunk is too large
            confirmed_breaks.append(candidate)
            last_break = candidate + 1
    
    return confirmed_breaks

def create_smart_chunks(sentences, breakpoints, min_chunk_size, max_chunk_size):
    """Create chunks with intelligent merging and splitting"""
    if not sentences:
        return []
    
    chunks = []
    start = 0
    
    # Create initial chunks based on breakpoints
    for bp in breakpoints:
        end = bp + 1
        chunk_sentences = sentences[start:end]
        if chunk_sentences:
            chunks.append(chunk_sentences)
        start = end
    
    # Add final chunk
    if start < len(sentences):
        final_chunk = sentences[start:]
        if final_chunk:
            chunks.append(final_chunk)
    
    # Post-process chunks for optimal sizing
    optimized_chunks = []
    i = 0
    
    while i < len(chunks):
        current_chunk = chunks[i]
        current_size = sum(s['char_length'] for s in current_chunk)
        
        # If chunk is too small, try to merge with next
        if current_size < min_chunk_size and i + 1 < len(chunks):
            next_chunk = chunks[i + 1]
            next_size = sum(s['char_length'] for s in next_chunk)
            
            if current_size + next_size <= max_chunk_size:
                # Merge chunks
                merged_chunk = current_chunk + next_chunk
                optimized_chunks.append(merged_chunk)
                i += 2  # Skip next chunk as it's merged
                continue
        
        # If chunk is too large, split it
        if current_size > max_chunk_size:
            split_chunks = split_large_chunk(current_chunk, min_chunk_size, max_chunk_size)
            optimized_chunks.extend(split_chunks)
        else:
            optimized_chunks.append(current_chunk)
        
        i += 1
    
    return optimized_chunks if optimized_chunks else chunks

def split_large_chunk(chunk, min_chunk_size, max_chunk_size):
    """Split a large chunk into smaller, coherent pieces"""
    if not chunk:
        return []
    
    total_chars = sum(s['char_length'] for s in chunk)
    target_chunks = max(2, total_chars // max_chunk_size + 1)
    target_size = total_chars // target_chunks
    
    split_chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in chunk:
        sentence_size = sentence['char_length']
        
        # If adding this sentence would exceed target size and we have content
        if current_size + sentence_size > target_size and current_chunk:
            # Only split if current chunk meets minimum size
            if current_size >= min_chunk_size:
                split_chunks.append(current_chunk)
                current_chunk = [sentence]
                current_size = sentence_size
            else:
                # Add to current chunk anyway to meet minimum
                current_chunk.append(sentence)
                current_size += sentence_size
        else:
            current_chunk.append(sentence)
            current_size += sentence_size
    
    # Add remaining chunk
    if current_chunk:
        split_chunks.append(current_chunk)
    
    return split_chunks

def enrich_chunks_with_images_semantic(chunks, image_map_path):
    """
    Advanced image enrichment with better association logic:
    1. Direct caption matching
    2. Reference-based association
    3. Proximity-based inclusion
    4. Context-aware grouping
    """
    try:
        with open(image_map_path, "r") as f:
            image_map = json.load(f)
        
        # Try to load detailed mapping
        detailed_map_path = image_map_path.replace('.json', '_detailed.json')
        if os.path.exists(detailed_map_path):
            with open(detailed_map_path, "r") as f:
                detailed_map = json.load(f)
        else:
            detailed_map = {k: {'file': v, 'position': 0, 'counter': i+1} 
                          for i, (k, v) in enumerate(image_map.items())}
            
    except (FileNotFoundError, json.JSONDecodeError):
        image_map = {}
        detailed_map = {}

    def extract_all_image_numbers(text):
        """Extract all possible image/figure references"""
        numbers = set()
        
        # Direct mentions
        patterns = [
            r"Image\s+(\d+)",
            r"Figure\s+(\d+)",
            r"image\s+(\d+)",
            r"figure\s+(\d+)"
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            numbers.update(int(n) for n in matches)
        
        # Contextual references
        ref_patterns = [
            r"(?:see|refer to|shown in|as depicted in|according to)\s+(?:image|figure)\s+(\d+)",
            r"(?:above|below|following)\s+(?:image|figure)\s+(\d+)"
        ]
        
        for pattern in ref_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            numbers.update(int(n) for n in matches)
        
        return list(numbers)

    def find_matching_images(image_numbers, image_map, detailed_map):
        """Find best matching images for given numbers"""
        matched_images = []
        
        for img_num in image_numbers:
            best_match = None
            best_score = 0
            
            for image_key, image_file in image_map.items():
                # Extract number from image key
                key_matches = re.findall(r"(?:Image|Figure)\s+(\d+)", image_key, re.IGNORECASE)
                if key_matches and int(key_matches[0]) == img_num:
                    score = 1.0
                    if score > best_score:
                        best_score = score
                        best_match = {
                            'label': image_key,
                            'file': image_file,
                            'number': img_num
                        }
                        if image_key in detailed_map:
                            best_match.update(detailed_map[image_key])
            
            if best_match:
                matched_images.append(best_match)
        
        return matched_images

    enriched_chunks = []
    used_images = set()
    
    for chunk_idx, chunk in enumerate(chunks):
        chunk_text = '\n'.join([s['sentence'] for s in chunk])
        
        # Find all image references in chunk
        image_numbers = extract_all_image_numbers(chunk_text)
        
        # Get matching images
        chunk_images = find_matching_images(image_numbers, image_map, detailed_map)
        
        # Check for standalone captions that should include their images
        for sentence in chunk:
            if sentence['is_caption']:
                caption_label = extract_label(sentence['sentence'])
                if caption_label and caption_label in image_map:
                    # Add the image if not already included
                    if not any(img['label'] == caption_label for img in chunk_images):
                        chunk_images.append({
                            'label': caption_label,
                            'file': image_map[caption_label],
                            'number': 0  # Will be determined from label
                        })
        
        # Remove duplicates and sort by number/position
        unique_images = []
        seen_files = set()
        
        for img in chunk_images:
            if img['file'] not in seen_files:
                unique_images.append(img)
                seen_files.add(img['file'])
                used_images.add(img['label'])
        
        # Sort by image number for consistent ordering
        unique_images.sort(key=lambda x: x.get('number', 999))
        
        # Prepare output format
        image_labels = [img['label'] for img in unique_images]
        image_files = [{'label': img['label'], 'file': img['file']} for img in unique_images]
        
        enriched_chunks.append({
            "chunk_id": chunk_idx,
            "chunk_text": chunk_text,
            "sentence_count": len(chunk),
            "char_count": len(chunk_text),
            "word_count": len(chunk_text.split()),
            "image_labels": image_labels,
            "image_files": image_files,
            "has_images": len(image_files) > 0,
            "has_captions": any(s['is_caption'] for s in chunk),
            "paragraph_span": {
                "start": min(s['paragraph_index'] for s in chunk),
                "end": max(s['paragraph_index'] for s in chunk)
            }
        })

    # Report on unused images
    all_image_labels = set(image_map.keys())
    unused_images = all_image_labels - used_images
    if unused_images:
        print(f"Warning: {len(unused_images)} images not associated with any chunk:")
        for label in sorted(unused_images):
            print(f"  - {label}")

    return enriched_chunks

def load_or_generate_enriched_chunks_semantic():
    """
    Load or generate advanced semantically chunked content with images
    """
    docx_mtime = get_file_modified_time(DOCX_LOCAL_PATH)
    map_mtime = get_file_modified_time(IMAGE_MAP_PATH)
    chunks_mtime = get_file_modified_time(ENRICHED_CHUNKS_PATH)

    if (chunks_mtime and docx_mtime and map_mtime and 
        chunks_mtime > docx_mtime and chunks_mtime > map_mtime):
        # Use cached version
        try:
            with open(ENRICHED_CHUNKS_PATH, "r") as f:
                chunks = json.load(f)
                print(f"Loaded {len(chunks)} cached semantic chunks")
                return chunks
        except (FileNotFoundError, json.JSONDecodeError):
            pass
    
    # Ensure image extraction is done first
    if not os.path.exists(IMAGE_MAP_PATH):
        print("Extracting images from DOCX...")
        extract_images_and_labels_from_docx(DOCX_LOCAL_PATH, IMAGE_DIR, IMAGE_MAP_PATH, debug=True)
    
    # Generate advanced semantic chunks
    print("Creating advanced semantic chunks...")
    chunks = semantic_chunking_docx(
        DOCX_LOCAL_PATH,
        model_name='all-MiniLM-L6-v2',
        buffer_size=2,
        percentile=85,
        min_chunk_size=150,
        max_chunk_size=1500,
        similarity_threshold=0.3
    )
    
    # Enrich with advanced image information
    print("Enriching chunks with advanced image association...")
    enriched_chunks = enrich_chunks_with_images_semantic(chunks, IMAGE_MAP_PATH)
    
    # Save to disk
    os.makedirs(os.path.dirname(ENRICHED_CHUNKS_PATH), exist_ok=True)
    with open(ENRICHED_CHUNKS_PATH, "w") as f:
        json.dump(enriched_chunks, f, indent=2)
    
    print(f"Generated {len(enriched_chunks)} advanced semantic chunks")
    
    # Print summary statistics
    analyze_advanced_chunks(enriched_chunks)
    
    return enriched_chunks

def get_file_modified_time(filepath):
    if os.path.exists(filepath):
        return datetime.fromtimestamp(os.path.getmtime(filepath))
    return None

def analyze_advanced_chunks(chunks):
    """Enhanced analysis of generated chunks"""
    print(f"\n=== ADVANCED CHUNK ANALYSIS ===")
    print(f"Total chunks: {len(chunks)}")
    
    # Statistics
    char_counts = [chunk['char_count'] for chunk in chunks]
    word_counts = [chunk['word_count'] for chunk in chunks]
    image_counts = [len(chunk['image_files']) for chunk in chunks]
    
    print(f"\nSize Statistics:")
    print(f"  - Avg characters: {np.mean(char_counts):.0f} (min: {min(char_counts)}, max: {max(char_counts)})")
    print(f"  - Avg words: {np.mean(word_counts):.0f} (min: {min(word_counts)}, max: {max(word_counts)})")
    print(f"  - Chunks with images: {sum(1 for c in chunks if c['has_images'])}/{len(chunks)}")
    print(f"  - Chunks with captions: {sum(1 for c in chunks if c['has_captions'])}/{len(chunks)}")
    
    # Image distribution
    print(f"\nImage Distribution:")
    for i in range(max(image_counts) + 1):
        count = sum(1 for c in image_counts if c == i)
        print(f"  - {count} chunks with {i} images")
    
    # Sample chunks
    print(f"\nSample Chunks:")
    for i, chunk in enumerate(chunks[:3]):
        print(f"\nChunk {i+1}:")
        print(f"  - Length: {chunk['char_count']} chars, {chunk['word_count']} words")
        print(f"  - Images: {len(chunk['image_files'])}")
        if chunk['image_files']:
            for img in chunk['image_files']:
                print(f"    * {img['label']} -> {img['file']}")
        print(f"  - Preview: {chunk['chunk_text'][:150]}...")
        print(f"  - Paragraph range: {chunk['paragraph_span']['start']}-{chunk['paragraph_span']['end']}")
